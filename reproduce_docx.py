from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def set_styles(doc):
    # Set Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    # Also need to set rFonts for complex scripts if needed, but simple name usually works for core fonts
    style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman') # Assuming this is desired

    # Set Heading 1
    style = doc.styles['Heading 1']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(15)
    style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # Set Heading 2 (Assuming 13 based on "everything else 13")
    style = doc.styles['Heading 2']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def main():
    doc = Document()
    set_styles(doc)

    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("1. First Section", level=1)
    doc.add_paragraph("Content of first section.")

    doc.add_heading("1.1. Subsection", level=2)
    doc.add_paragraph("Content of subsection.")

    doc.add_heading("2. Second Section", level=1)
    doc.add_paragraph("Content of second section.")

    doc.save("test_output.docx")
    print("Created test_output.docx")

if __name__ == "__main__":
    main()
