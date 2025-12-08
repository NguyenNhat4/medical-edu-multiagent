from docx import Document

def generate_doc(content_data, filename, topic="Medical Document"):
    """
    Generates a Word document from a list of content data.

    Args:
        content_data (list): List of dicts, each containing 'title', 'content', 'speaker_notes'.
        filename (str): Output filename.
        topic (str): Main title of the document.
    """
    doc = Document()
    doc.add_heading(topic, 0)

    for item in content_data:
        title = item.get("title", "Untitled")
        doc.add_heading(title, level=1)

        # Content bullets
        points = item.get("content", [])
        if isinstance(points, str):
            doc.add_paragraph(points)
        elif isinstance(points, list):
            for p in points:
                doc.add_paragraph(p, style='List Bullet')

        # Notes as separate paragraph or subsection
        notes = item.get("speaker_notes", "")
        if notes:
            doc.add_heading("Notes / Clinical Relevance:", level=2)
            doc.add_paragraph(notes)

    doc.save(filename)
    return filename
